import os
import datetime
import zlib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Path import *
from Doc import *
from icecream import ic
from base64 import urlsafe_b64encode as b64e, urlsafe_b64decode as b64d


def obscure(unobscured: str) -> str:
    return b64e(zlib.compress(str.encode(unobscured), 9)).decode()


def unobscure(obscured: str) -> str:
    return zlib.decompress(b64d(str.encode(obscured))).decode()


def write_invitation(doc, components):
    date_on_document = datetime.datetime.strptime(components['date on letter'].get(), "%b %d, %Y")
    has_second_host = False

    bearer_of_expenses = {
        'host(s)': "paid for by me and will be my responsibility",
        'guest(s)': "their own responsibility and will be paid for by themselves. I will provide additional support if any assistance is needed",
    }

    document_data = {
        '[DAY]': date_on_document.strftime("%d"),
        '[MONTH]': date_on_document.strftime("%b"),
        '[YEAR]': date_on_document.strftime("%Y"),
        '[BEARER]': bearer_of_expenses.get(components["bearer of expenses"].get(), bearer_of_expenses['host(s)']),
        '[GUEST_PURPOSE]': components['purpose of visit'].get(),
        '[GUEST_ARRIVAL]': components['arrival date'].get(),
        '[GUEST_DEPARTURE]': components['departure date'].get(),
        '[GUEST_CANADIAN_ADDRESS]': components['address in Canada'].get(),
        '[GUEST_RESIDENCE]': components['country of residence'].get(),
        '[GUEST_NAMES]': [],
        '[GUEST_INFO]': [],
    }

    for i in range(1,3):
        if components[f'host {i} name'].get().strip() != '':
            other_host = "2" if i==1 else "1"
            document_data[f'[HOST{i}_NAME]'] = components[f"host {i} name"].get().strip()
            document_data[f'[HOST{i}_BIRTH]'] = components[f"host {i} date of birth"].get().strip()
            document_data[f'[HOST{i}_PASSPORT]'] = components[f"host {i} passport no."].get().strip()
            document_data[f'[HOST{i}_ADDRESS]'] = components[f"host {i} address"].get().strip()
            document_data[f'[HOST{i}_PHONE]'] = components[f"host {i} phone number"].get().strip()
            document_data[f'[HOST{i}_EMAIL]'] = components[f"host {i} email address"].get().strip()
            document_data[f'[HOST{i}_OCCUPATION]'] = components[f"host {i} occupation"].get().strip()
            document_data[f'[HOST{i}_STATUS]'] = components[f"host {i} status in Canada"].get().strip()
            document_data[f'[HOST{i}_RELATION_TO_HOST{other_host}]'] = components[f'relationship to host {other_host}'].get().lower().strip()

        has_second_host = '[HOST2_NAME]' in document_data

    for i in range(1,6):
        if components[f'guest {i} name'].get().strip() != '':
            document_data['[GUEST_NAMES]'].append(components[f'guest {i} name'].get().strip())
            document_data['[GUEST_INFO]'].append(f"{components[f'guest {i} name'].get().strip()}, born {components[f'guest {i} date of birth'].get().strip()}, passport #{components[f'guest {i} passport no.'].get().strip()}")
            document_data[f'[GUEST{i}_NAME]'] = components[f"guest {i} name"].get().strip()
            document_data[f'[GUEST{i}_BIRTH]'] = components[f"guest {i} date of birth"].get().strip()
            document_data[f'[GUEST{i}_PASSPORT]'] = components[f"guest {i} passport no."].get().strip()
            document_data[f'[GUEST{i}_ADDRESS]'] = components[f"guest {i} address"].get().strip()
            document_data[f'[GUEST{i}_PHONE]'] = components[f"guest {i} phone number"].get().strip()
            document_data[f'[GUEST{i}_EMAIL]'] = components[f"guest {i} email address"].get().strip()
            document_data[f'[GUEST{i}_OCCUPATION]'] = components[f"guest {i} occupation"].get().strip()
            document_data[f'[GUEST{i}_CITIZENSHIP]'] = components[f"guest {i} country of citizenship"].get().strip()
            document_data[f'[GUEST{i}_RELATION_TO_HOST1]'] = components[f"guest {i} relation to host 1"].get().strip()

    def replace_last(string, delimiter, replacement):
        start, _, end = string.rpartition(delimiter)
        return start + replacement + end

    document_data['[GUEST_NAMES]'] = replace_last((', ').join(document_data['[GUEST_NAMES]']), (", "), ", and ")
    document_data['[GUEST_INFO]'] = ('\n').join(document_data['[GUEST_INFO]'])

    if has_second_host:
        doc = Document(resource_path("assets\\templates\\invitation_2.docx"))
        intro_text = (f"This letter is to express me and my {document_data[f'[HOST1_RELATION_TO_HOST2]'].lower()}'s interest in inviting {document_data['[GUEST_NAMES]']} to Canada and to furthermore support their Temporary Resident Visa application.")
    else:
        intro_text = (f"This letter is to express my interest in inviting {document_data['[GUEST_NAMES]']} to Canada and to furthermore support their Temporary Resident Visa application.")

    outro_text = (f"\n\n{components['conclusion content'].get()}\nIf any clarification or information is required, please do not hesitate to contact us at our email addresses and phone numbers below.\n\n")

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = PT(10)

    overwrite_placeholders(doc, document_data)
    insert_paragraph(paragraph=doc.add_paragraph(), text=intro_text) # add intro

    for host_number in range(1,3):
        if f'[HOST{host_number}_NAME]' in document_data:
            insert_2col_table(
                document=doc, 
                table_heading="\nMy details are as follows," if host_number == 1 else f"\n\nMy {document_data.get('[HOST2_RELATION_TO_HOST1]', '')}'s details are as follows,", 
                table_items=[
                    {"label": "Full Name", "info": document_data[f'[HOST{host_number}_NAME]']},
                    {"label": "Date of Birth", "info": document_data[f'[HOST{host_number}_BIRTH]']},
                    {"label": "Canadian Status", "info": document_data[f'[HOST{host_number}_STATUS]']},
                    {"label": "Passport Number", "info": document_data[f'[HOST{host_number}_PASSPORT]']},
                    {"label": "Residential Address", "info": document_data[f'[HOST{host_number}_ADDRESS]']},
                    {"label": "Phone Number", "info": document_data[f'[HOST{host_number}_PHONE]']},
                    {"label": "Email Address", "info": document_data[f'[HOST{host_number}_EMAIL]']},
                    {"label": "Current Occupation", "info": document_data[f'[HOST{host_number}_OCCUPATION]']},
                ]
            )

    if has_second_host:
        doc.add_page_break()

    for guest_number in range(1,6):
        if f'[GUEST{guest_number}_NAME]' in document_data:
            insert_2col_table(
                document=doc, 
                table_heading=f"\n\n{document_data[f'[GUEST{guest_number}_NAME]']}'s details are as follows,", 
                table_items=[
                    {"label": "Full Name", "info": document_data[f'[GUEST{guest_number}_NAME]']},
                    {"label": "Date of Birth", "info": document_data[f'[GUEST{guest_number}_BIRTH]']},
                    {"label": "Citizen of", "info": document_data[f'[GUEST{guest_number}_CITIZENSHIP]']},
                    {"label": "Passport Number", "info": document_data[f'[GUEST{guest_number}_PASSPORT]']},
                    {"label": "Residential Address", "info": document_data[f'[GUEST{guest_number}_ADDRESS]']},
                    {"label": "Phone Number", "info": document_data[f'[GUEST{guest_number}_PHONE]']},
                    {"label": "Email Address", "info": document_data[f'[GUEST{guest_number}_EMAIL]']},
                    {"label": "Current Occupation", "info": document_data[f'[GUEST{guest_number}_OCCUPATION]']},
                    {"label": "Relationship to Inviter", "info": document_data[f'[GUEST{guest_number}_RELATION_TO_HOST1]']},
                    {"label": "Arrival Date", "info": document_data[f'[GUEST_ARRIVAL]']},
                    {"label": "Departure Date", "info": document_data[f'[GUEST_DEPARTURE]']},
                    {"label": "Primary Residence in Canada", "info": document_data[f'[GUEST_CANADIAN_ADDRESS]']},
                ]
            )

    insert_paragraph(paragraph=doc.add_paragraph(), text=outro_text) # add conclusion

    # table for hosts to sign
    client_signatures = [
        {
            "label_l": f"{document_data.get('[HOST1_NAME]', '')}\n{document_data.get('[HOST1_EMAIL]', '')}\n{document_data.get('[HOST1_PHONE]', '')}",
            "info_l": "",
            "label_r": f"{document_data.get('[HOST2_NAME]', '')}\n{document_data.get('[HOST2_EMAIL]', '')}\n{document_data.get('[HOST2_PHONE]', '')}",
            "info_r": "",
        },
    ]

    insert_4col_table(document=doc, table_heading="", table_items=client_signatures)

    timestamp_obj = datetime.datetime.now()
    formatted_timestamp = str(timestamp_obj.strftime("%Y.%m.%d-%H.%M.%S"))
    output_filename = f"[{formatted_timestamp}] - {components[f"host 1 name"].get().strip()} - Invitation Letter"
    save_doc(doc=doc, components=components, override_output_filename=output_filename, folder_name='invitations')


def overwrite_placeholders(doc, document_data):
    try:
        for paragraph in doc.paragraphs:
            for key, value in document_data.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(key, value)

    except Exception as e:
        print(e)


def get_prompt_response(prompt="") -> str:
    import google.generativeai as genai
    from dotenv import load_dotenv

    if prompt.strip() == "":
        return ""

    load_dotenv()
    genai.configure(api_key=os.getenv('API_KEY'))

    # See https://ai.google.dev/api/python/google/generativeai/GenerativeModel
    generation_config = {
        "temperature": 1,
        "top_p": 0.95,
        "top_k": 64,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }

    model = genai.GenerativeModel(
        model_name="gemini-2.5-flash",
        generation_config=generation_config,
    )

    chat_session = model.start_chat(history=[])
    response = chat_session.send_message(prompt)

    return(response.text)

