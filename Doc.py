import datetime, os, re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from GUI import *
from docx.shared import Cm as CM
from docx.shared import Pt as PT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# format date to `{date + suffix} {full month name} {year}`
def format_date(date_string):
    if date_string == 'advance':
        return date_string

    temp = datetime.datetime.strptime(date_string, '%d/%m/%Y')
    return str(format_day(temp.strftime("%d")) + " " + temp.strftime("%B") + " " + temp.strftime("%Y"))


# add suffix 'th'/'st'/'rd' to days
def format_day(day):
    day = int(day)
    suffix = ""

    if (4 <= day <= 20 or 24 <= day <= 30):
        suffix = "th"
    else:
        suffix = ["st", "nd", "rd"][day % 10 - 1]

    return str(day) + suffix


# format phone number to add brackets and hyphens when applicable
def format_phone(number):
    if (len(number) == 10):
        area = number[0] + number[1] + number[2]
        prefix = number[3] + number[4] + number[5]
        line = number[6] + number[7] + number[8] + number[9]
        return "(" + area + ")" + " " + prefix + "-" + line

    return number


# 
def set_cell_color(cell, color="ffffff"):
    cell_xml = cell._tc
    cell_props = cell_xml.get_or_add_tcPr()
    shade_obj = OxmlElement("w:shd")
    shade_obj.set(qn("w:fill"), color)
    cell_props.append(shade_obj)


# 
def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


# 
def set_row_color(row, color="FFFFFF"):
    for cell in row.cells:
        set_cell_color(cell, color)


# 
def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existance, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


#
def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


#
def insert_paragraph(paragraph, text=None):
    paragraph.insert_paragraph_before(text)


# set up folders and save files, print if needed
def save_doc(doc=None, components=None, folder_name="", prefix="", override_output_filename="") -> bool:
    if doc is None or components is None and override_output_filename.strip() == "":
        return False

    try:
        # set up the output directory
        output_dir = f'{os.getcwd()}\\output\\{folder_name.lower()}\\'.replace('\\\\', '\\')

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        if len(override_output_filename.strip()) > 0:
            output_filename = override_output_filename
        else:
            output_filename = f"{prefix} - {re.sub(r' \((.*?)\)', '', components['client 1 first name'].get()).strip()} {re.sub(r' \((.*?)\)', '', components['client 1 last name'].get()).strip()}"

        output_filename = f"{output_filename}{'.docx' if '.docx' not in output_filename else ''}"

        # save the file to the output folder
        doc.save(f"{output_dir}{output_filename}")

        # open the word file
        # os.startfile(output_dir + output_filename)

    except Exception as e:
        # ErrorPopup(f"Exception when saving and opening document:\n\n{str(e)}")
        print(f"Exception when saving document: {e}")
        return False

    return True


# write the invoice number and date on the top
def insert_invoice_info(document=None, doc_id="", billed_to="", timestamp=""):

    if (len(billed_to.strip()) == 0):
        insert_paragraph(
            document.add_paragraph(),
            (f"Payment#\t{doc_id}\nDate\t\t{timestamp}\n\n"),
        )
    else:
        insert_paragraph(
            document.add_paragraph(),
            (f"Payment#\t{doc_id}\nDate\t\t{timestamp}\nBilled to\t{billed_to}\n"),
        )


# itemized table
def insert_items_table(doc=None, cart_items=[]):

    # Table data in a form of list
    headings = [
        {"heading": "SL.", "width": 1.0},
        {"heading": "DESCRIPTION", "width": 20.0},
        {"heading": "QTY", "width": 1.0},
        {"heading": "RATE", "width": 1.0},
        {"heading": "AMOUNT", "width": 1.0},
    ]

    # Creating a table object
    items_table = doc.add_table(rows=1, cols=5)

    # add heading in the 1st row of the table
    row = items_table.rows[0].cells
    for idx, col in enumerate(headings):
        row[idx].text = col["heading"]

    # add data from the list to the table
    for index, entry in enumerate(cart_items):

        # Adding a row and then adding data in it.
        row = items_table.add_row().cells

        row[0].text = str(entry["serial"])
        row[1].text = str(entry["service"])
        row[2].text = str(entry["quantity"])
        row[3].text = "{:,.2f}".format(entry["rate"])
        row[4].text = "{:,.2f}".format(int(entry["quantity"]) * float(entry["rate"]))

    # set the table borders
    for cell in items_table.rows[0].cells:
        set_cell_border(cell, bottom={"sz": 6, "color": "#EEE", "val": "single", "space": "10"})
        set_cell_border(cell, top={"sz": 6, "color": "#EEE", "val": "single", "space": "15"})

    # set column widths
    for idx, col in enumerate(headings):
        for index, cell in enumerate(items_table.columns[idx].cells):
            cell.width = CM(col["width"])

            if index > 0:
                set_cell_border(cell, bottom={"sz": 6, "color": "#EEE", "val": "single", "space": "15"},)

        set_cell_border(cell, bottom={"sz": 6, "color": "#EEE", "val": "single", "space": "0"})

    # set row heights
    for index, row in enumerate(items_table.rows):
        row.height = CM(1)

    make_rows_bold(items_table.rows[0])


# table containing taxes and total
def insert_totals_table(doc, cart_items): 
    total_table = doc.add_table(rows=1, cols=5)

    taxes = 0
    total = 0
    for item in cart_items:
        taxes += float(item.get('taxes'))
        total += float(item.get('price'))

    total_table_data = [
        [str("TAXES"), "", "", "", "${:,.2f}".format(taxes)],
        [str("TOTAL"), "", "", "", "${:,.2f}".format(total)],
    ]

    for total_row in total_table_data:
        row = total_table.add_row().cells

        for i, col in enumerate(total_row):
            row[i].text = col

        make_rows_bold(total_table.rows[-1])


# insert table with the passed data
def insert_2col_table(document, table_heading=None, table_items=[], table_props=None):
    style = document.styles['Normal']
    font = style.font
    # font.name = 'Times New Roman'

    if table_props is None:
        table_props = {
            "color": "#AAAAAA",
            "cell_alignment": [
                WD_ALIGN_PARAGRAPH.LEFT, 
                WD_ALIGN_PARAGRAPH.RIGHT
            ]
        }

    # add the table heading
    try:
        if table_heading is not None:
            insert_paragraph(
                paragraph=document.add_paragraph(),
                text=table_heading,
                bolded=True,
                italicized=False
            )
    except Exception as e:
        ErrorPopup(f'Exception when adding table heading: {table_heading}\n\n{e}')
        return False

    # add the table contents
    try:
        # Creating a table object
        table_obj = document.add_table(rows=len(table_items), cols=len(table_items[0]))

        # add rows
        for idx, row in enumerate(table_obj.rows):
            curr_row = row.cells
            row.height = CM(0.50)
            curr_row[0].text = table_items[idx]['label']
            curr_row[0].paragraphs[0].alignment = table_props['cell_alignment'][0]
            curr_row[1].text = table_items[idx]['info']
            curr_row[1].paragraphs[0].alignment = table_props['cell_alignment'][0]

        # set column widths and borders
        for idx, _ in enumerate(table_items[0]):
            for cell in (table_obj.columns[0].cells):
                cell = table_obj.columns[0].cells[0]
                cell.width = CM(0)

            for cell in (table_obj.columns[1].cells):
                cell.width = CM(25)
                set_cell_border(cell, top={"sz": 1, "color": "#ffffff", "val": "single", "space": "8"})
                set_cell_border(cell, bottom={"sz": 1, "color": "#888888", "val": "single", "space": "0"})

    except Exception as e:
        ErrorPopup(f'Exception when adding table contents\n\n{e}')
        return False


# insert table with the passed data
def insert_4col_table(document, table_heading=None, table_items=[], table_props=None):
    style = document.styles['Normal']
    font = style.font
    # font.name = 'Times New Roman'

    if table_props is None:
        table_props = {
            "color": "#AAAAAA",
            "cell_alignment": [
                WD_ALIGN_PARAGRAPH.LEFT, 
                WD_ALIGN_PARAGRAPH.RIGHT
            ]
        }

    # add the table heading
    try:
        if table_heading is not None:
            insert_paragraph(
                paragraph=document.add_paragraph(),
                text=table_heading,
                bolded=True,
                italicized=False
            )
    except Exception as e:
        # ErrorPopup(f'Exception when adding table heading: {table_heading}\n\n{e}')
        print(f'Exception when adding table heading: {table_heading}\n\n{e}')
        return False

    # add the table contents
    try:
        # Creating a table object
        table_obj = document.add_table(rows=len(table_items), cols=len(table_items[0]))

        # add rows
        for idx, row in enumerate(table_obj.rows):
            curr_row = row.cells
            row.height = CM(0.50)
            curr_row[0].text = table_items[idx]['label_l']
            curr_row[0].paragraphs[0].alignment = table_props['cell_alignment'][0]
            curr_row[1].text = table_items[idx]['info_l']
            curr_row[1].paragraphs[0].alignment = table_props['cell_alignment'][0]
            curr_row[2].text = table_items[idx]['label_r']
            curr_row[2].paragraphs[0].alignment = table_props['cell_alignment'][0]
            curr_row[3].text = table_items[idx]['info_r']
            curr_row[3].paragraphs[0].alignment = table_props['cell_alignment'][0]

        # set column widths and borders
        for idx, _ in enumerate(table_items[0]):

            for cell in (table_obj.columns[0].cells):
                cell = table_obj.columns[0].cells[0]
                cell.width = CM(8)

            for cell in (table_obj.columns[1].cells):
                cell.width = CM(14)
                set_cell_border(cell, top={"sz": 1, "color": "#ffffff", "val": "single", "space": "8"})
                if len(cell.text) > 0:
                    set_cell_border(cell, bottom={"sz": 1, "color": "#888888", "val": "single", "space": "0"})

            for cell in (table_obj.columns[2].cells):
                cell.width = CM(8)

            for cell in (table_obj.columns[3].cells):
                cell.width = CM(14)
                set_cell_border(cell, top={"sz": 1, "color": "#ffffff", "val": "single", "space": "4"})
                if len(cell.text) > 0:
                    set_cell_border(cell, bottom={"sz": 1, "color": "#888888", "val": "single", "space": "0"})

    except Exception as e:
        # ErrorPopup(f'Exception when adding table contents\n\n{e}')
        print(f'Exception when adding table contents\n\n{e}')
        return False


# insert a new paragraph
def insert_paragraph(paragraph, text="", bolded=False, italicized=False):

    if bolded or italicized:
        runner = paragraph.add_run(text)
        runner.bold = bolded
        runner.italic = italicized

    else:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        paragraph.paragraph_format.space_before = PT(0)
        paragraph.paragraph_format.space_after = PT(0)
        new_para = Paragraph(new_p, paragraph._parent)
        new_para.add_run(text)


