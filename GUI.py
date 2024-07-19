import math
import customtkinter as ctk
import datetime
import os
from Popups import ErrorPopup
import actions
from Path import *
from docx import Document
from CTkMessagebox import CTkMessagebox
from tkinter import StringVar
from icecream import ic
from subprocess import DEVNULL, STDOUT, check_call
from writer import write_auth, replace_placeholders, write_receipt
from dateutil import relativedelta as rd
from typing import Literal
from actions import HistoryViewer, ReceiptFinder, add_item_button, remove_item_button, test_button, decrypt_button


family_medium="Roboto Bold"
family_bold="Roboto Bold"


class GUI:
    def __init__(self, master=None, label_text="", left_offset=0, top_offset=0) -> None:
        # no label text was passed
        if label_text == "":
            label_text = f"Component_{datetime.datetime.now().timestamp()}"

        self.stringvar = StringVar(value="")
        self.component = None

        self.label = ctk.CTkLabel(master, width=190, text=label_text, anchor="w", font=ctk.CTkFont(family=family_bold))
        self.label.grid(row=top_offset, column=0, pady=10, padx=5, columnspan=1)

    def get(self) -> str:
        field = self.component
        return field.get()

    def set(self, new_text: str = "") -> None:
        self.stringvar.set(new_text)

    def reset(self) -> None:
        self.stringvar.set("")


class RowBreak(GUI):
    def __init__(self, master=None, left_offset=0, top_offset=0, heading="") -> None:
        self.breakline = ctk.CTkLabel(master, text=heading.upper(), height=32, width=450, fg_color="#808080", text_color="white", corner_radius=2, font=ctk.CTkFont(family=family_bold, weight='bold'))
        self.breakline.grid(row=top_offset, column=0, pady=10, padx=5, columnspan=5)

    def reset(self) -> None:
        return


class ComboBox(GUI):
    def __init__(self, master=None, app=None, label_text="", options=None, left_offset=0, top_offset=0, default_string='click to select', default_option=None) -> None:
        """create a new GUI ComboBox object"""

        super().__init__(master, label_text, left_offset, top_offset)

        self.options = options
        self.default_string = default_string
        self.stringvar = StringVar(value="no options added" if options is None else self.default_string)
        self.default_option = None

        self.component = ctk.CTkComboBox(
            master,
            width=250,
            height=32,
            border_width=0,
            corner_radius=2,
            bg_color="#fff",
            fg_color="#ddd",
            values=options,
            variable=self.stringvar,
            font=ctk.CTkFont(family=family_medium),
            dropdown_font=ctk.CTkFont(family=family_medium),
        )

        if default_option is not None:
            self.default_option = default_option
            self.set(default_option)

        # self.component.place(x=left_offset + 210, y=top_offset + 8)
        self.component.grid(row=top_offset, column=1, pady=10, padx=5, columnspan=3)

    def dropdown_callback(self, choice):
        self.callback(self.app_components)

    def get(self) -> str:
        """returns the first option if nothing was selected"""
        if self.component.get() == self.default_string:
            return self.options[0]
        else:
            return self.component.get()

    def set(self, opt) -> None:
        self.stringvar.set(opt)

    def reset(self) -> None:
        if self.default_option is None:
            self.component.set(self.default_string)
        else:
            self.component.set(self.default_option)

    def add_callback(self, component_name="", app=None, callback=None):
        if callback is None or app is None:
            return

        self.app_components = app.get_all_components()
        self.callback = callback
        self.component.configure(command=self.dropdown_callback)


class Entry(GUI):
    def __init__(self, master=None, app=None, label_text="", left_offset=0, top_offset=0, placeholder="", default_text=None) -> None:
        """create a new GUI Entry object"""

        super().__init__(master, label_text, left_offset, top_offset)

        self.stringvar = StringVar(value="")
        self.default_text=default_text

        self.component = ctk.CTkEntry(
            master,
            width=250,
            height=32,
            border_width=0,
            corner_radius=2,
            bg_color="#fff",
            fg_color="#ddd",
            textvariable=self.stringvar,
            font=ctk.CTkFont(family=family_medium),
        )

        self.reset()
        self.component.grid(row=top_offset, column=1, pady=10, padx=5, columnspan=3)

    def reset(self) -> None:
        self.stringvar.set(self.default_text if self.default_text is not None else '')


    def stringvar_callback(self, *args):
        self.callback(self.app_components)

    def add_callback(self, component_name="", app=None, callback=None):
        if callback is None or app is None:
            return

        self.app_components = app.get_all_components()
        self.callback = callback
        self.stringvar.trace_add('write', self.stringvar_callback)


class DatePicker(GUI):
    def __init__(self, master=None, label_text="", left_offset=0, top_offset=0, show_day=True, populate_years_with=[]) -> None:
        """create a new GUI DatePicker object"""

        super().__init__(master, label_text, left_offset, top_offset)

        self.today = datetime.datetime.now()
        self.show_day = show_day

        self.stringvar_month = StringVar(value=self.today.strftime("%b"))
        self.stringvar_day = StringVar(value=self.today.strftime("%d"))
        self.stringvar_year = StringVar(value=self.today.strftime("%Y"))

        self.component_day = ctk.CTkComboBox(
            master,
            width=70,
            height=32,
            border_width=0,
            corner_radius=2,
            bg_color="#fff",
            fg_color="#ddd",
            values=self.populate_days(),
            variable=self.stringvar_day,
            font=ctk.CTkFont(family=family_medium)
        )

        # in some cases, like credit card expirations, the day is not needed
        if show_day is True:
            self.component_day.grid(row=top_offset, column=1, pady=10, padx=5)
        else:
            ctk.CTkLabel(master, height=32, width=70, text="").grid(row=top_offset, column=1, pady=10, padx=5)

        self.component_month = ctk.CTkComboBox(
            master,
            width=80,
            height=32,
            border_width=0,
            corner_radius=2,
            bg_color="#fff",
            fg_color="#ddd",
            values=self.populate_months(),
            variable=self.stringvar_month,
            command=self.repopulate_days,
            font=ctk.CTkFont(family=family_medium),
        )

        self.component_month.grid(row=top_offset, column=2, pady=10, padx=5)

        self.component_year = ctk.CTkComboBox(
            master,
            width=80,
            height=32,
            border_width=0,
            corner_radius=2,
            bg_color="#fff",
            fg_color="#ddd",
            values=self.populate_years(populate_years_with),
            variable=self.stringvar_year,
            command=self.repopulate_days,
            font=ctk.CTkFont(family=family_medium),
        )

        self.component_year.grid(row=top_offset, column=3, pady=10, padx=5)


    # returns a list of days depending on the month
    def populate_days(self) -> list:
        days=[]

        months = {
            "Jan": "31",
            "Feb": "29" if (int(self.stringvar_year.get()) % 4 == 0) else "28",
            "Mar": "31",
            "Apr": "30",
            "May": "31",
            "Jun": "30",
            "Jul": "31",
            "Aug": "31",
            "Sep": "30",
            "Oct": "31",
            "Nov": "30",
            "Dec": "31",
        }

        selected_month = months[(self.stringvar_month.get())]

        for i in range(1, int(selected_month)+1):
            days.append(str(i))

        return days


    # returns a list of month names
    def populate_months(self) -> list:
        return ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]


    # returns a list of years
    def populate_years(self, populate_years_with) -> list:
        if len(populate_years_with) > 0:
            return populate_years_with

        years = []

        for i in range(10):
            next_year = int(self.stringvar_year.get()) - 10 + i
            years.append(str(next_year))

        for i in range(11):
            next_year = int(self.stringvar_year.get()) + i
            years.append(str(next_year))

        return years


    # recalculates the number of days to pick from, based on the month
    def repopulate_days(self, _) -> None:
        self.component_day.configure(values=self.populate_days())


    # set the date picker back to the current date
    def reset(self) -> None:
        self.stringvar_month.set(self.today.strftime("%b"))
        self.stringvar_day.set(self.today.strftime("%d"))
        self.stringvar_year.set(self.today.strftime("%Y"))


    # return a formatted date
    def get(self) -> str:
        m = self.stringvar_month.get()
        d = self.stringvar_day.get()
        y = self.stringvar_year.get()

        if self.show_day is True:
            return f"{m} {d}, {y}"

        return f"{m}, {y}"


    # set the date 
    def set(self, m: str|int = None, d: str|int = None, y: str|int = None) -> str:
        if m is None:
            m = self.today.strftime("%b")
        if d is None:
            d = self.today.strftime("%d")
        if y is None:
            y = self.today.strftime("%Y")

        if type(m) is str:
            self.stringvar_month.set(m)
        else:
            monthnames = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
            self.stringvar_month.set(monthnames[m])

        self.stringvar_day.set(str(d))
        self.stringvar_year.set(str(y))

        return self.get()


class PaymentSplitter(GUI):
    def __init__(self, master=None, app=None, label_text="", left_offset=0, top_offset=0) -> None:

        super().__init__(master, label_text, left_offset, top_offset)

        self.app = app

        self.pay_amount = Entry(master=master, label_text=label_text, left_offset=10, top_offset=top_offset)
        self.pay_amount.component.configure(width=152)
        self.pay_amount.stringvar.set(value="$")
        self.pay_amount.component.grid(row=top_offset, column=1, pady=10, padx=5, columnspan=2)

        self.split_quantity = ComboBox(
            master=master, 
            label_text=label_text, 
            left_offset=10, 
            top_offset=top_offset, 
            default_string = "number of months", 
            options=["1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12"],
        )

        self.split_quantity.component.configure(width=172, command=self.dropdown_callback)
        self.split_quantity.component.grid(row=top_offset, column=3, pady=10, padx=5, columnspan=2)

        # shorten the width of the label to fit the window
        self.label.configure(width=100, text=label_text)

        # get rid of the labels that come with the Entry and DatePicker objects
        self.pay_amount.label.destroy()
        self.split_quantity.label.destroy()


    def reset(self) -> None:
        self.pay_amount.set("$")
        self.split_quantity.set("number of months")


    def get(self, part=None) -> dict[str, str]:
        payment = {
            "amount": float(0) if self.pay_amount.get().replace("$", "") == "" else float(self.pay_amount.get().replace("$", "")),
            "months": int(self.split_quantity.get())
        }

        if part == "amount":
            return payment['amount']
        if part == "months":
            return payment['months']

        return payment


    def set(self, amount, opt) -> None:
        self.pay_amount.set(amount)
        self.split_quantity.set(opt)


    def dropdown_callback(self, choice):
        self.split_payment()


    def split_payment(self) -> None:
        amount = 0 if self.pay_amount.get() == "" else float(self.pay_amount.get().replace("$", ""))
        months = int(self.split_quantity.get())
        amount_per_month = float(amount/months)
        self.pay_amount.set("${:.2f}".format(float(self.pay_amount.get().replace("$", ""))))


        # if there is nothing to pay each month, simply reset each payment widget
        if (amount_per_month == 0):
            components = self.app.get_all_components()

            for i in range(12):
                components.get(f'payment {i+1}').reset()

            return

        # set the amount to be paid each month and change up the colors a bit
        curr_month = 0
        start_point = ""
        for component_name, component_obj in zip(self.app.components.keys(), self.app.components.values()):

            if "payment 1" in component_name:
                start_point = component_obj.get('date')

            if "payment" in component_name:
                dt_object = datetime.datetime.strptime(start_point, "%b %d, %Y")
                dt_object = dt_object + rd.relativedelta(days=curr_month*30)

                component_obj.set(
                    "{:.2f}".format(amount_per_month), 
                    month=dt_object.strftime("%b"), 
                    date=dt_object.strftime("%d"), 
                    year=dt_object.strftime("%Y")
                )

                component_obj.label.configure(text_color="#000")
                component_obj.pay_amount.component.configure(fg_color="light green", text_color="#000")
                component_obj.pay_date.component_day.configure(fg_color="light green", text_color="#000")
                component_obj.pay_date.component_month.configure(fg_color="light green", text_color="#000")
                component_obj.pay_date.component_year.configure(fg_color="light green", text_color="#000")

                curr_month += 1

            if months < curr_month:
                component_obj.reset()
                component_obj.label.configure(text_color="#bbb")
                component_obj.pay_amount.component.configure(fg_color="#ddd", text_color="#aaa")
                component_obj.pay_date.component_day.configure(fg_color="#ddd", text_color="#aaa")
                component_obj.pay_date.component_month.configure(fg_color="#ddd", text_color="#aaa")
                component_obj.pay_date.component_year.configure(fg_color="#ddd", text_color="#aaa")

            if curr_month == 12:
                break


class PaymentInfo(GUI):
    def __init__(self, master=None, label_text="", left_offset=0, top_offset=0) -> None:

        super().__init__(master, label_text, left_offset, top_offset)

        self.pay_amount = Entry(master=master, label_text=label_text, left_offset=10, top_offset=top_offset)
        self.pay_amount.component.configure(width=70, fg_color="#ddd")
        self.pay_amount.stringvar.set(value="$")
        self.pay_amount.component.grid(row=top_offset, column=1, pady=10, padx=5, columnspan=1)

        dt_object_now = datetime.datetime.strftime(datetime.datetime.now(), "%Y")
        dt_object_max = datetime.datetime.strftime(datetime.datetime.now() + rd.relativedelta(months=12), "%Y")

        self.pay_date = DatePicker(master=master, label_text=label_text, left_offset=10, top_offset=top_offset, populate_years_with=[dt_object_now, dt_object_max])
        self.pay_date.component_day.grid(row=top_offset, column=2, pady=10, padx=5)
        self.pay_date.component_month.grid(row=top_offset, column=3, pady=10, padx=5)
        self.pay_date.component_year.grid(row=top_offset, column=4, pady=10, padx=5)

        # shorten the width of the label to fit the window
        self.label.configure(width=100, text=label_text)

        # get rid of the labels that come with the Entry and DatePicker objects
        self.pay_amount.label.destroy()
        self.pay_date.label.destroy()


    def reset(self) -> None:
        self.label.configure(text_color="black")
        self.pay_amount.component.configure(fg_color="#ddd", text_color="#000")
        self.pay_date.component_day.configure(fg_color="#ddd", text_color="#000")
        self.pay_date.component_month.configure(fg_color="#ddd", text_color="#000")
        self.pay_date.component_year.configure(fg_color="#ddd", text_color="#000")
        self.pay_amount.stringvar.set("$")
        self.pay_date.reset()


    def get(self, part=None) -> dict[float, int]:
        payment = {
            "amount": float(0) if self.pay_amount.get().replace("$", "") == "" else float(self.pay_amount.get().replace("$", "")),
            "date": self.pay_date.get()
        }

        if part == "amount":
            return payment['amount']
        if part == "date":
            return payment['date']

        return payment


    def set(self, amount, year, month, date) -> None:
        self.pay_amount.set(f"${amount}")
        self.pay_date.set(y=year, m=month, d=date)


class AppButton():
    def __init__(self, app=None, master=None, image=None, left_offset=0, top_offset=0, app_name="", width=72, height=72, desc="", row=0) -> None:

        self.component = ctk.CTkButton(
            master=master,
            text="",
            image=image,
            border_width=0,
            corner_radius=2,
            fg_color="transparent",
            command=lambda:self.__open_app(app_name=app_name, app=app),
            width=width,
            height=height,
        ).grid(row=row, column=0, pady=10, padx=5)

        self.desc_frame = ctk.CTkFrame(
            master=master, 
            corner_radius=8, 
            border_width=1, 
            width=256, 
            height=68, 
            fg_color="#ffffff",
        )

        self.desc_frame.grid(row=row, column=1, pady=10, padx=5)

        self.desc_text = ctk.CTkLabel(master=self.desc_frame, text=app_name, width=240, wraplength=256, anchor="w", font=ctk.CTkFont(family=family_bold, weight='bold')).place(x=10, y=8)
        self.desc_text = ctk.CTkLabel(master=self.desc_frame, text=desc, width=240, wraplength=256, anchor="w", font=ctk.CTkFont(family=family_medium), text_color="#777777").place(x=10, y=32)


    def __open_app(self, app_name, app):
        owd = os.getcwd()

        try:
            os.chdir(f"{os.getcwd()}\\assets\\apps\\{app_name}\\")
            app.hide()
            check_call([f"{app_name}.exe"])
        except Exception as e:
            ErrorPopup(msg=f"Could not find {app_name} in path:\n\n{os.getcwd()}\\assets\\apps\\{app_name}\\")

        app.unhide()
        os.chdir(owd)


class ActionButton():
    def __init__(self, app=None, action="", master=None, image=None, btn_text="", btn_color="transparent", width=91, height=40, row=0, col=0, blueprint={}) -> None:

        self.component = ctk.CTkButton(
            master=master,
            text=btn_text,
            image=image,
            border_width=0,
            corner_radius=2,
            fg_color=btn_color,
            command=lambda:self.assign_action(app, action, blueprint),
            width=width,
            height=height,
            state='disabled' if 'spacer' in action else 'normal'
        ).grid(row=row, column=col, pady=10, padx=5)


    def assign_action(self, app=None, action="", blueprint={}) -> None:
        if ("reset" in action):
            app_components = app.get_all_components()
            for component_name in blueprint.keys():
                if (component_name in app_components):
                    app_components.get(component_name).reset()

            if ("receipt" in action):
                app.components.get('cart').tools.buttons[3].configure(fg_color='light gray', text_color="white", state='disabled', text="$0.00")
                app.components.get('cart').tools.buttons[4].configure(fg_color='light gray', text_color="white", state='disabled', text="$0.00")

        elif (action == "payments"):
            try:
                doc = Document(resource_path("assets\\templates\\payments.docx"))
                write_auth(doc, app.get_all_components())
            except Exception as e:
                ErrorPopup(msg=f'Exception while writing payment authorization:\n\n{str(e)}')

        elif (action == "retainer"):
            try:
                doc = Document(resource_path("assets\\templates\\retainer.docx"))
                replace_placeholders(doc, app.get_all_components(), "Retainer")
            except Exception as e:
                ErrorPopup(msg=f'Exception while writing retainer:\n\n{str(e)}')

        elif (action == "conduct"):
            try:
                doc = Document(resource_path("assets\\templates\\conduct.docx"))
                replace_placeholders(doc, app.get_all_components(), "Conduct")
            except Exception as e:
                ErrorPopup(msg=f'Exception while writing code of conduct:\n\n{str(e)}')

        elif (action == "create receipt"):
            try:
                doc = Document(resource_path("assets\\templates\\receipt.docx"))
                write_receipt(doc, app.get_all_components())
            except Exception as e:
                ErrorPopup(msg=f'Exception while writing receipt:\n\n{str(e)}')

        elif (action == "history"):
            HistoryViewer(app)

        elif (action == "find receipt"):
            ReceiptFinder(app)

        elif (action == "test"):
            test_button(app)

        elif (action == "decrypt"):
            decrypt_button(app)

        elif (action == "add item"):
            add_item_button(app)

        elif (action == "remove item"):
            remove_item_button(app)

        elif (action == "output"):
            try:
                os.startfile(os.getcwd() + "\\output")
            except Exception as e:
                ErrorPopup(msg=f'Output folder not found')


class TabView():
    def __init__(self, master, new_tabs=[]) -> None:

        self.TabView = ctk.CTkTabView(master, corner_radius=2, fg_color="#fff")
        self.TabView.pack(expand=True, fill="both", padx=10, pady=10)

        if len(new_tabs) > 0:
            self.tabs = {}
            self.set_tabs(new_tabs)

    def set_tabs(self, new_tabs):
        for tab in new_tabs:
            tab_obj = self.TabView.add(tab)
            self.tabs[tab] = tab_obj

    def get_tabs(self):
        tabs = {}

        for tab_name in self.tabs:
            tabs[tab_name] = self.TabView.tab(name=tab_name) 

        return tabs


class WindowView():
    def __init__(self, window_name="_window_", app=None, width=0, height=0) -> None:
        self.body = ctk.CTkToplevel()

        w = app.get_size('w')*0.5 if width == 0 else width
        h = app.get_size('h')*0.5 if height == 0 else height
        x = (app.get_size('w')/2) - (w/2)
        y = (app.get_size('h')/2) - (h/2)

        self.body.title(window_name)
        self.body.geometry('%dx%d+%d+%d' % (w, h, x, y))
        self.body.resizable(False, False)
        self.body.configure(fg_color='white')
        self.body.after(300, lambda:self.show())

    def show(self) -> None:
        self.body.focus()


class CellWidget():
    def __init__(self, master=None, width=0, height=0, text="", text_color="black", fg_color="#fff", info=[], font=None, type='label', command=None, hover_color=None, on_enter=None, on_leave=None) -> None:
        self.info = info
        self.width = width
        self.height = height
        self.text = text

        if type == 'label':
            self.cell = ctk.CTkLabel(
                master=master, 
                width=width, 
                height=height, 
                text=text, 
                text_color=text_color, 
                fg_color=fg_color, 
                font=font,
            )

        elif type == 'button':
            self.cell = ctk.CTkButton(
                master=master, 
                width=width, 
                height=height, 
                text=text, 
                text_color=text_color, 
                fg_color=fg_color, 
                font=font,
                corner_radius=0,
                command=command,
                hover_color=hover_color,
            )

            if on_enter is not None or on_leave is not None:
                self.set_hover_methods(None, on_enter, on_leave)


    def set_hover_methods(self, hover_color=None, on_enter=None, on_leave=None):
        if hover_color is not None:
            self.cell.configure(hover_color=hover_color)

        self.cell.bind('<Enter>', on_enter, add='+')
        self.cell.bind('<Leave>', on_leave, add='+')

    def get_cell(self):
        return self.cell

    def get_info(self):
        return self.info

    def destroy(self):
        self.cell.destroy()


class RowWidget():
    def __init__(self, parent_frame=None, row_contents=[], row_color="#eee", row_content_methods=[None, None, None], parent_width=0, row_number=0, mode:Literal["header", "tools", "table"]="table", is_blank = False):

        # no parent, nowhere to put this
        if parent_frame is None:
            return
        
        if len(row_contents) == 0:
            if is_blank:
                row_contents=['','','','','']
            else:
                row_contents=["col_0", "col_1", "col_2", "col_3", "col_4"]

        self.container = ctk.CTkFrame(master=parent_frame, fg_color="white", border_width=0, width=parent_width, height=30)
        self.buttons = []
        self.contents = []
        self.info = {'selected': False}

        def highlight_row():
            for c in self.contents:
                c.cell.configure(fg_color='#7ac8ff')

        def unhighlight_row():
            for c in self.contents:
                c.cell.configure(fg_color=row_color)

        def select_row():
            self.info['selected'] = True
            ic(self.info)

        # setup the grid system
        for i in range(len(row_contents)+1):
            parent_frame.columnconfigure(index=i, weight=1)

        # set the number of buttons based on the mode
        # 5 in a tools row where all 5 cols are buttons
        for i in range(5 if mode=="tools" else 0):
            self.buttons.append(
                ctk.CTkButton(
                    master=self.container,
                    text=row_contents[i],
                    text_color="white", 
                    border_width=0,
                    corner_radius=0,
                    fg_color="black" if row_contents[i] != "" else "light gray",
                    width=(parent_width-61)/5,
                    height=38,
                    font=ctk.CTkFont(family=family_bold, size=12, weight='bold'),
                    state="disabled" if row_contents[i] == "" else "normal",
                    command=row_content_methods[i],
                )
            )

            self.buttons[i].grid(row=0, column=i, pady=0, padx=0)

        # row_contents are placed as CTkLabel widgets only in table and header modes
        # tools mode does not place any CTkLabel widgets
        if mode != "tools":
            for i, content in enumerate(row_contents):
                self.contents.append(
                    CellWidget(
                        master=self.container, 
                        type='button' if (not is_blank and i == 0 and mode != 'header') else 'label',
                        width=(parent_width-61)/5, 
                        height=38, 
                        text=content, 
                        text_color="white" if mode is "header" else "black", 
                        fg_color="#000" if mode is "header" else row_color, 
                        font=ctk.CTkFont(family=family_bold, size=12, weight='bold') if mode is "header" else ctk.CTkFont(family=family_medium, size=12),
                        command=lambda: select_row(),
                        on_enter=lambda *args: highlight_row(),
                        on_leave=lambda *args: unhighlight_row(),
                    )
                )

                self.contents[i].get_cell().grid(row=0, column=i + (0 if mode=="header" else 1), pady=0, padx=0)

        # place the entire container with all the stuff above
        self.container.grid(row=row_number, column=0, pady=0, padx=0)

    def set(self, row_contents=["col_0", "col_1", "col_2", "col_3", "col_4"], row_info=None):
        for index, col in enumerate(self.contents):
            col.get_cell().configure(text=row_contents[index])

        if row_info is None:
            return

        elif len(row_info) <= 0:
            self.info = {}
        else:
            self.info = row_info

    def get_cells(self) -> list:
        row_contents = []

        for col in self.contents:
            row_contents.append(col)

        return row_contents

    def get_info(self) -> list:
        return self.info

    def cleanup(self):
        for widget in self.buttons:
            widget.destroy()

        for widget in self.contents:
            widget.destroy()

        self.container.destroy()


class TableWidget():
    def __init__(self, master=None, headers:list=[], rows:list=[], parent_width=0, parent_height=0, rows_per_page=15, active_row=0):

        self.rows = rows
        self.parent_frame = master
        self.parent_width = parent_width
        self.parent_height = parent_height
        self.rows_per_page = rows_per_page
        self.rows_rendered = []
        self.active_row = active_row
        self.page = 1

        for i in range(3):
            master.columnconfigure(index=i, weight=1)

        self.header_frame = ctk.CTkFrame(
            master=self.parent_frame, 
            fg_color="white", 
            border_width=0, 
            width=self.parent_width, 
            height=self.parent_height*0.05, 
        )

        self.header = RowWidget(
            parent_frame=self.header_frame, 
            parent_width=parent_width, 
            mode="header", 
            row_contents=headers, 
        )

        self.table_frame = ctk.CTkFrame(
            master=self.parent_frame, 
            fg_color="white", 
            border_width=1, 
            width=self.parent_width, 
            height=self.parent_height*0.90, 
        )

        self.tools_frame = ctk.CTkFrame(
            master=self.parent_frame, 
            fg_color="white", 
            border_width=0, 
            width=self.parent_width, 
            height=self.parent_height*0.05
        )

        self.tools = RowWidget(
            parent_frame=self.tools_frame, 
            parent_width=self.parent_width, 
            mode="tools", 
            row_contents=[
                "previous page", 
                "next page", 
                "", 
                "", 
                ""
            ], 
            row_content_methods=[
                lambda:self.navigate(page=self.page-1),
                lambda:self.navigate(page=self.page+1),
                lambda:None,
                lambda:None,
                lambda:None,
            ])

        self.header_frame.grid(row=0, column=1, pady=[9,0])
        self.table_frame.grid(row=1, column=1, pady=[0,0])
        self.tools_frame.grid(row=2, column=1, pady=[2,0])

        self.reset()


    def set_custom_method(self, index, func):
        self.tools.buttons[index+2].configure(command=func)


    def navigate(self, page=0):
        if (page == 1):
            self.tools.buttons[0].configure(fg_color="light gray", state="disabled")
            self.tools.buttons[1].configure(fg_color="black", state="normal")

        elif (page == math.ceil(len(self.rows)/self.rows_per_page)-1):
            self.tools.buttons[0].configure(fg_color="black", state="normal")
            self.tools.buttons[1].configure(fg_color="light gray", state="disabled")

        else:
            self.tools.buttons[0].configure(fg_color="black", state="normal")
            self.tools.buttons[1].configure(fg_color="black", state="normal")

        for row in self.rows_rendered:
            row.cleanup()

        self.page=page
        self.refresh()
        self.update(page=page)


    def refresh(self):

        for r in self.rows_rendered:
            r.cleanup()

        self.table_frame.destroy()
        self.table_frame = ctk.CTkFrame(
            master=self.parent_frame, 
            fg_color="white", 
            border_width=1, 
            width=self.parent_width, 
            height=self.parent_height*0.90, 
        )

        self.table_frame.grid(row=1, column=1, pady=[2,0])


    def update(self, page=1):

        self.refresh()
        page_offset = self.rows_per_page * (page-1)

        for index in range(self.rows_per_page):
            if (index + page_offset) < len(self.rows):
                self.rows_rendered.append(
                    RowWidget(
                        parent_frame=self.table_frame, 
                        parent_width=self.rows[index + page_offset].get('parent_width'), 
                        row_number=self.rows[index + page_offset].get('row_number'), 
                        mode=self.rows[index + page_offset].get('mode'), 
                        row_contents=self.rows[index + page_offset].get('row_contents'),
                        row_color=self.rows[index + page_offset].get('row_color'),
                    )
                )

                # set the next button to be active if the next row is not a blank
                if (index + page_offset) == len(self.rows) - 1:
                    self.tools.buttons[1].configure(fg_color="light gray", state="disabled")
                else:
                    self.tools.buttons[1].configure(fg_color="black", state="normal")

            else:
                RowWidget(
                    parent_frame=self.table_frame, 
                    parent_width=self.parent_width, 
                    row_number=index + page_offset, 
                    row_color="#ddd" if ((index + page_offset) % 2 == 0) else "#eee",
                    is_blank = True,
                    mode='table', 
                )

                # set the next button to be active if the last row was blank
                self.tools.buttons[1].configure(fg_color="light gray", state="disabled")


    def reset(self) -> None:
        for row in self.rows_rendered:
            row.cleanup()

        self.table_frame.destroy()
        self.table_frame = ctk.CTkFrame(
            master=self.parent_frame, 
            fg_color="white", 
            border_width=1, 
            width=self.parent_width, 
            height=self.parent_height*0.90, 
        )

        self.table_frame.grid(row=1, column=1, pady=[2,0])

        self.rows = []
        self.rows_rendered = []
        self.active_row = 0
        self.page = 1

        for i in range(self.rows_per_page):
            RowWidget(
                parent_frame=self.table_frame,
                parent_width=self.parent_width,
                is_blank=True,
                row_color="#ddd" if i % 2==0 else "#eee",
                row_number=i,
                mode="table",
            )

        self.tools.buttons[0].configure(fg_color="light gray", state="disabled")

        if len(self.rows) <= self.rows_per_page:
            self.tools.buttons[1].configure(fg_color="light gray", state="disabled")


    def add(self, row_info=None, row_contents=None, row_index=None) -> None:
        for row_info, row_contents in zip(row_info, row_contents):

            row_to_update = row_index if row_index is not None else self.active_row

            new_row = (
                {
                    'parent_frame': self.table_frame, 
                    'parent_width': self.parent_width, 
                    'row_number': row_to_update, 
                    'mode': "table", 
                    'row_contents': row_contents if not None else ['','','','',''],
                    'row_color': "#ddd" if (row_to_update % 2 == 0) else "#eee",
                    'info': row_info,
                }
            )

            if row_index is None:
                self.rows.append(new_row)
                self.active_row += 1
            else:
                self.rows[row_index] = new_row

        self.update()
        self.update(page=self.page)


    def remove(self) -> dict:
        rows_after_removal = []
        ic(self.rows[0])

        # for row in self.rows[0]:
        #     if 'selected' in row.info:
        #         ic(row)

        #     if row.get('selected') is False:
        #         rows_after_removal.append(row)

        # self.rows = rows_after_removal
        # ic(self.rows)


    def contains(self, row_info=[], compare_keys=[]):
        for row_index in range(self.active_row):

            current_row_match = False

            for k in compare_keys:
                if self.rows[row_index].get('info').get(k) != row_info[k]:
                    current_row_match = False
                    break
                else:
                    current_row_match = True

            if (current_row_match):
                return (row_index, self.rows[row_index].get('info'))

        return (None, None)


    def get(self) -> list:
        return self.rows

